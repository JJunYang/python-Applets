import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

students = pd.read_excel('students.xlsx')
students.sort_values(by='Score', inplace=True, ascending=False)

plt.bar(students['Name'], students.Score, color='orange')

plt.title('Student Score', fontsize=16)
plt.xlabel('Name')
plt.ylabel('score')

plt.xticks(students.Name, rotation='90')
plt.tight_layout()
# plt.show()
imgname = 'data.jpg'
plt.savefig(imgname)

document=Document()
document.add_heading('Data Analysis',level=0)
first_student=students.iloc[0,:]['Name']
first_student_score=students.iloc[0,:]['Score']
p=document.add_paragraph('No1 Student is:')
p.add_run(str(first_student)).bold=True
p.add_run(', score:')
p.add_run(str(first_student_score)).bold=True

p1=document.add_paragraph(f'Total {len(students.Name)} students')

table=document.add_table(rows=len(students.Name)+1,cols=2)

table.style='LightShading-Accent1'

table.cell(0,0).text='Name'
table.cell(0,1).text='Score'

for i,(index,row) in enumerate(students.iterrows()):
    table.cell(i+1,0).text=str(row['Name'])
    table.cell(i+1,1).text=str(row['Score'])

document.add_picture(imgname)
document.save('Students.docx')


