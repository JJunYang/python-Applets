from docx import Document
from docx.shared import Inches

document=Document()

p1=document.add_paragraph('first paragraph')

p2=document.add_paragraph('second paragraph')

p3=p2.insert_paragraph_before('insert paragraph')

document.add_heading('heading',level=1)

document.add_page_break()

table=document.add_table(rows=5,cols=5)

cell=table.cell(0,3)
cell.text='first row,fourth col'
row=table.rows[1]
row.cells[0].text='second row,first col'
row.cells[1].text='second row,second col'

document.add_picture('img.jpeg',width=Inches(1.25))
document.save('new.docx')
